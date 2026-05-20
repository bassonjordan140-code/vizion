# -*- coding: utf-8 -*-
"""
ViZion — Générateur de Rapport d'Audit Énergétique
===================================================
Usage :
    python generate_report.py <dossier_rapport>

Exemple :
    python generate_report.py rapports/2026-05-20_14h30_Hotel-Exemple

Le script lit audit.json + photos/ dans le dossier et génère un fichier
rapport_audit.docx dans ce même dossier.
"""

import json
import os
import sys
import io
from datetime import datetime

# Fix encodage Windows pour les print
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
#   COULEURS DU RAPPORT (identiques au modèle TIPEE)
# ============================================================

TEAL = RGBColor(0x17, 0x88, 0x88)       # #178888 — titres, en-têtes
DARK_TEAL = RGBColor(0x10, 0x60, 0x60)  # sous-titres
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_BG = RGBColor(0xE8, 0xF4, 0xF4)  # fond léger tableau

FONT_NAME = "Calibri"

# ============================================================
#   HELPERS
# ============================================================

def set_cell_shading(cell, color_hex):
    """Applique un fond coloré à une cellule de tableau."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style_name=None, bold=False, size=11,
                         color=BLACK, alignment=None, space_after=6, space_before=0):
    """Ajoute un paragraphe stylé."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_heading_teal(doc, text, level=1):
    """Ajoute un titre avec la couleur teal."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL
        run.font.name = FONT_NAME
    return h


def add_info_table(doc, rows_data, col_widths=None):
    """Crée un tableau clé/valeur (2 colonnes)."""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for label, value in rows_data:
        row = table.add_row()
        c0 = row.cells[0]
        c1 = row.cells[1]

        c0.text = str(label)
        c1.text = str(value) if value else "—"

        # Style label (bold, fond teal léger)
        set_cell_shading(c0, "E8F4F4")
        for p in c0.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME

        for p in c1.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = FONT_NAME

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph()  # espace
    return table


def add_multi_col_table(doc, headers, rows_data):
    """Crée un tableau multi-colonnes avec en-tête teal."""
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # En-tête
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_shading(cell, "178888")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME
                run.font.color.rgb = WHITE

    # Lignes
    for row_data in rows_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val) if val else "—"
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_NAME

    doc.add_paragraph()
    return table


def try_add_photo(doc, photos_dir, photo_key, caption="", width_inches=3.5):
    """Tente d'insérer une photo depuis le dossier photos/."""
    if not photos_dir or not os.path.isdir(photos_dir):
        return False

    # Chercher le fichier qui contient la clé
    for fname in os.listdir(photos_dir):
        if photo_key in fname:
            fpath = os.path.join(photos_dir, fname)
            try:
                doc.add_picture(fpath, width=Inches(width_inches))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    add_styled_paragraph(doc, caption, size=9, color=GRAY,
                                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
                return True
            except Exception as e:
                print(f"  ⚠ Photo ignorée ({fname}): {e}")
                return False
    return False


def oui_non(val):
    """Convertit un booléen ou 'oui'/'non' en texte."""
    if isinstance(val, bool):
        return "Oui" if val else "Non"
    if isinstance(val, str):
        return "Oui" if val.lower() == "oui" else "Non"
    return "—"


def val_or_dash(val, suffix=""):
    """Retourne la valeur ou '—'."""
    if val is None or val == "" or val == 0:
        return "—"
    return f"{val}{suffix}"


# ============================================================
#   PAGE DE GARDE
# ============================================================

def build_cover_page(doc, audit_data):
    """Construit la page de garde du rapport."""
    meta = audit_data.get("meta", {})

    # Espace haut
    for _ in range(3):
        doc.add_paragraph()

    # Bandeau titre
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "178888")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAPPORT")
    run.font.size = Pt(36)
    run.font.color.rgb = WHITE
    run.font.name = FONT_NAME
    run.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("AUDIT ÉNERGÉTIQUE")
    run2.font.size = Pt(36)
    run2.font.color.rgb = WHITE
    run2.font.name = FONT_NAME
    run2.bold = True

    doc.add_paragraph()

    # Nom du site
    site_name = "Établissement"
    donnees = audit_data.get("donnees", {})
    # Chercher un nom dans hébergement
    if "hebergements" in donnees:
        heb = donnees["hebergements"]
        first_key = list(heb.keys())[0] if heb else None
        if first_key and heb[first_key].get("nom"):
            site_name = heb[first_key]["nom"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(site_name.upper())
    run.font.size = Pt(30)
    run.font.color.rgb = TEAL
    run.font.name = FONT_NAME
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Infos
    date_str = meta.get("exportDateFormatted", datetime.now().strftime("%d/%m/%Y %H:%M"))
    add_info_table(doc, [
        ("Date d'export", date_str),
        ("Version application", f"ViZion v{meta.get('appVersion', '0.1')}"),
        ("Nombre de photos", str(meta.get("nbPhotos", 0))),
    ], col_widths=[6, 10])

    doc.add_page_break()


# ============================================================
#   SECTION 1 — MODULES SÉLECTIONNÉS
# ============================================================

def build_modules_overview(doc, audit_data):
    """Vue d'ensemble des modules audités."""
    add_heading_teal(doc, "1. MODULES AUDITÉS", level=1)

    modules = audit_data.get("modulesSelectionnes", [])
    if not modules:
        add_styled_paragraph(doc, "Aucun module sélectionné.")
        return

    rows = []
    for mod in modules:
        rows.append([mod.get("label", mod.get("id", "")), str(mod.get("quantity", 1))])

    add_multi_col_table(doc, ["Module", "Nombre de fiches"], rows)
    doc.add_page_break()


# ============================================================
#   SECTION 2 — ÉTAT DES LIEUX
# ============================================================

def build_etat_des_lieux(doc, audit_data, photos_dir):
    """Section principale avec toutes les données terrain."""
    add_heading_teal(doc, "2. ÉTAT DES LIEUX", level=1)

    donnees = audit_data.get("donnees", {})
    photo_index = audit_data.get("photosIndex", [])

    # --- HÉBERGEMENTS ---
    if "hebergements" in donnees:
        add_heading_teal(doc, "2.1 Hébergements", level=2)
        for num, data in donnees["hebergements"].items():
            build_hebergement(doc, num, data, photos_dir)

    # --- PISCINES ---
    if "piscines" in donnees:
        add_heading_teal(doc, "2.2 Piscines", level=2)
        for num, data in donnees["piscines"].items():
            build_piscine(doc, num, data, photos_dir)

    # --- RESTAURANT ---
    if "restaurant" in donnees:
        add_heading_teal(doc, "2.3 Restaurants", level=2)
        for num, data in donnees["restaurant"].items():
            build_restaurant(doc, num, data, photos_dir)

    # --- BAR ---
    if "bar" in donnees:
        add_heading_teal(doc, "2.4 Bars", level=2)
        for num, data in donnees["bar"].items():
            build_bar(doc, num, data, photos_dir)

    # --- SPA ---
    if "spa" in donnees:
        add_heading_teal(doc, "2.5 Spas", level=2)
        for num, data in donnees["spa"].items():
            build_spa(doc, num, data, photos_dir)

    # --- BUANDERIE ---
    if "buanderie" in donnees:
        add_heading_teal(doc, "2.6 Buanderies", level=2)
        for num, data in donnees["buanderie"].items():
            build_buanderie(doc, num, data, photos_dir)

    # --- CUISINE ---
    if "cuisine" in donnees:
        add_heading_teal(doc, "2.7 Cuisines", level=2)
        for num, data in donnees["cuisine"].items():
            build_cuisine(doc, num, data, photos_dir)

    # --- BUREAUX ---
    if "bureaux" in donnees:
        add_heading_teal(doc, "2.8 Bureaux", level=2)
        for num, data in donnees["bureaux"].items():
            build_bureaux(doc, num, data, photos_dir)

    # --- SALLE DE RÉUNION ---
    if "reunion" in donnees:
        add_heading_teal(doc, "2.9 Salles de réunion", level=2)
        for num, data in donnees["reunion"].items():
            build_reunion(doc, num, data, photos_dir)

    # --- SALLE DE SPORT ---
    if "sport" in donnees:
        add_heading_teal(doc, "2.10 Salles de sport", level=2)
        for num, data in donnees["sport"].items():
            build_sport(doc, num, data, photos_dir)

    # --- SALLE DE JEUX ---
    if "jeux" in donnees:
        add_heading_teal(doc, "2.11 Salles de jeux", level=2)
        for num, data in donnees["jeux"].items():
            build_jeux(doc, num, data, photos_dir)

    # --- PARKING ---
    if "parking" in donnees:
        add_heading_teal(doc, "2.12 Parkings", level=2)
        for num, data in donnees["parking"].items():
            build_parking(doc, num, data, photos_dir)


# ============================================================
#   BUILDERS PAR MODULE
# ============================================================

def build_climatisation_section(doc, data, module_prefix, num, photos_dir):
    """Section climatisation commune à plusieurs modules."""
    clim = data.get("climatisation", {})
    if not clim or not clim.get("present"):
        add_styled_paragraph(doc, "Climatisation : Non")
        return

    add_info_table(doc, [
        ("Climatisation présente", "Oui"),
        ("Nombre d'unités", val_or_dash(clim.get("nombre"))),
        ("État", clim.get("etat", "—")),
        ("Plaque signalétique", oui_non(clim.get("plaque", False))),
    ])

    try_add_photo(doc, photos_dir, f"{module_prefix}_{num}_clim_plaque",
                  "Photo plaque signalétique climatisation")
    try_add_photo(doc, photos_dir, f"{module_prefix}_{num}_clim",
                  "Photo climatisation")


def build_brasseur_section(doc, data):
    """Section brasseur d'air commune."""
    brasseur = data.get("brasseurAir", {})
    present = brasseur.get("present")
    if isinstance(present, str):
        present = present == "oui"
    if not present:
        return
    add_info_table(doc, [
        ("Brasseur d'air", "Oui"),
        ("Nombre", val_or_dash(brasseur.get("nombre"))),
    ])


def build_eclairage_section(doc, data, module_prefix, num, photos_dir):
    """Section éclairage commune."""
    eclairages = data.get("eclairages", [])
    if not eclairages:
        add_styled_paragraph(doc, "Aucun éclairage renseigné.")
        return

    rows = []
    total_w = 0
    for ecl in eclairages:
        p = ecl.get("puissance", 0)
        q = ecl.get("quantite", 0)
        t = p * q
        total_w += t
        rows.append([
            ecl.get("type", "—"),
            f"{p} W",
            str(q),
            f"{t} W"
        ])

    rows.append(["TOTAL", "", "", f"{total_w} W"])
    add_multi_col_table(doc, ["Type", "W unitaire", "Quantité", "Total"], rows)

    # Photos éclairage
    for i in range(len(eclairages)):
        try_add_photo(doc, photos_dir, f"{module_prefix}_{num}_eclairage_{i}",
                      f"Photo éclairage {i+1}")


def build_equipements_checkbox(doc, equip_data, equip_defs):
    """Affiche les équipements cochés (checkbox)."""
    if not equip_data:
        add_styled_paragraph(doc, "Aucun équipement renseigné.")
        return

    rows = []
    for eq_id, eq_vals in equip_data.items():
        label = eq_id  # fallback
        # Chercher le label dans les defs
        for d in equip_defs:
            if d["id"] == eq_id:
                label = d["label"]
                break

        nombre = eq_vals.get("nombre", "—")
        rows.append([label, val_or_dash(nombre)])

    if rows:
        add_multi_col_table(doc, ["Équipement", "Nombre"], rows)


# --- HÉBERGEMENT ---

def build_hebergement(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Hébergement {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Nombre d'étages", val_or_dash(data.get("nbEtages"))),
        ("Nombre d'hébergements", val_or_dash(data.get("nbHebergements"))),
        ("Nombre de chambres", val_or_dash(data.get("nbChambres"))),
        ("Capacité", val_or_dash(data.get("capacite"))),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
    ])

    # Photos façade
    try_add_photo(doc, photos_dir, f"hebergement_{num}_facade", "Photo façade")

    # Climatisation
    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "hebergement", num, photos_dir)
    build_brasseur_section(doc, data)

    # Parois
    parois = data.get("parois", [])
    if parois:
        add_heading_teal(doc, "Matériaux de construction — Parois", level=4)
        rows = []
        for p in parois:
            rows.append([
                p.get("orientation", "—"),
                p.get("type", "—"),
                val_or_dash(p.get("epaisseur"), " cm"),
                p.get("isolation", "—")
            ])
        add_multi_col_table(doc, ["Orientation", "Type", "Épaisseur", "Isolation"], rows)

    # Ouvrants
    ouvrants = data.get("ouvrants", [])
    if ouvrants:
        add_heading_teal(doc, "Menuiseries — Ouvrants", level=4)
        rows = []
        for o in ouvrants:
            rows.append([
                o.get("type", "—"),
                o.get("vitrage", "—"),
                val_or_dash(o.get("quantite")),
            ])
        add_multi_col_table(doc, ["Type", "Vitrage", "Quantité"], rows)

    # Protections solaires
    protections = data.get("protectionsSolaires", [])
    if protections:
        add_heading_teal(doc, "Protections solaires", level=4)
        rows = []
        for pr in protections:
            rows.append([
                pr.get("type", "—"),
                pr.get("etat", "—"),
                val_or_dash(pr.get("quantite")),
            ])
        add_multi_col_table(doc, ["Type", "État", "Quantité"], rows)

    # Équipements chambre
    equips = data.get("equipements", [])
    if equips:
        add_heading_teal(doc, "Équipements des chambres", level=4)
        rows = []
        for eq in equips:
            if isinstance(eq, dict):
                rows.append([eq.get("nom", "—"), val_or_dash(eq.get("puissance"), " W")])
            else:
                rows.append([str(eq), "—"])
        if rows:
            add_multi_col_table(doc, ["Équipement", "Puissance"], rows)

    # ECS
    ecs = data.get("eauChaudeSanitaire", data.get("ecs", {}))
    if ecs and isinstance(ecs, dict):
        add_heading_teal(doc, "Eau chaude sanitaire", level=4)
        add_info_table(doc, [
            ("Type de production", ecs.get("type", "—")),
            ("Nombre de douches", val_or_dash(ecs.get("nbDouches"))),
        ])

    # Éclairage
    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "hebergement", num, photos_dir)

    # Observations
    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- PISCINE ---

def build_piscine(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Piscine {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Type", data.get("type", "—")),
        ("Volume (m³)", val_or_dash(data.get("volume"), " m³")),
        ("Température consigne (°C)", val_or_dash(data.get("tempConsigne"), " °C")),
        ("Couverture isotherme", oui_non(data.get("couvertureIsotherme", False))),
        ("Horaires", f"{data.get('horDebut', '—')} → {data.get('horFin', '—')}"),
    ])

    try_add_photo(doc, photos_dir, f"piscine_{num}_vue", "Photo piscine")

    # Pompes
    pompes = data.get("pompes", [])
    if pompes:
        add_heading_teal(doc, "Pompes de filtration", level=4)
        rows = []
        for p in pompes:
            rows.append([
                val_or_dash(p.get("puissance"), " kW"),
                f"{p.get('debitM3h', '—')} m³/h",
                f"{p.get('dureeH', '—')} h/jour",
            ])
        add_multi_col_table(doc, ["Puissance", "Débit", "Durée/jour"], rows)

    # Traitement
    add_info_table(doc, [
        ("Traitement de l'eau", data.get("traitementEau", "—")),
    ])

    # Éclairage
    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "piscine", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- RESTAURANT ---

def build_restaurant(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Restaurant {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Nombre de couverts", val_or_dash(data.get("nbCouverts"))),
        ("Horaires", f"{data.get('horDebut', '—')} → {data.get('horFin', '—')}"),
    ])

    try_add_photo(doc, photos_dir, f"restaurant_{num}_salle", "Photo salle")

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "restaurant", num, photos_dir)
    build_brasseur_section(doc, data)

    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "restaurant", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- BAR ---

def build_bar(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Bar {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Horaires", f"{data.get('horDebut', '—')} → {data.get('horFin', '—')}"),
    ])

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "bar", num, photos_dir)
    build_brasseur_section(doc, data)

    # Équipements bar
    equips = data.get("equipements", {})
    if equips:
        add_heading_teal(doc, "Équipements", level=4)
        rows = []
        for eq_id, eq_data in equips.items():
            rows.append([eq_id, val_or_dash(eq_data.get("nombre") if isinstance(eq_data, dict) else eq_data)])
        add_multi_col_table(doc, ["Équipement", "Nombre"], rows)

    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "bar", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- SPA ---

def build_spa(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Spa {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Horaires", f"{data.get('horDebut', '—')} → {data.get('horFin', '—')}"),
    ])

    # Bains
    bains = data.get("bains", [])
    if bains:
        add_heading_teal(doc, "Bains / Jacuzzis", level=4)
        for i, b in enumerate(bains):
            add_info_table(doc, [
                (f"Bain {i+1}", ""),
                ("Volume (m³)", val_or_dash(b.get("volume"), " m³")),
                ("Température (°C)", val_or_dash(b.get("temperature"), " °C")),
            ])

    # Hammams
    hammams = data.get("hammams", [])
    if hammams:
        add_heading_teal(doc, "Hammams", level=4)
        for i, h in enumerate(hammams):
            add_info_table(doc, [
                (f"Hammam {i+1}", ""),
                ("Puissance (kW)", val_or_dash(h.get("puissance"), " kW")),
            ])

    # Saunas
    saunas = data.get("saunas", [])
    if saunas:
        add_heading_teal(doc, "Saunas", level=4)
        for i, s in enumerate(saunas):
            add_info_table(doc, [
                (f"Sauna {i+1}", ""),
                ("Puissance (kW)", val_or_dash(s.get("puissance"), " kW")),
            ])

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "spa", num, photos_dir)

    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "spa", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- BUANDERIE ---

def build_buanderie(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Buanderie {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
    ])

    # Lave-linge
    ll = data.get("laveLinge", {})
    if ll:
        add_heading_teal(doc, "Lave-linge", level=4)
        add_info_table(doc, [
            ("Nombre", val_or_dash(ll.get("nombre"))),
            ("Capacité (kg)", val_or_dash(ll.get("capaciteKg"), " kg")),
            ("Cycles/jour", val_or_dash(data.get("cyclesLaveLinge"))),
        ])

    # Sèche-linge
    sl = data.get("secheLinge", {})
    if sl:
        add_heading_teal(doc, "Sèche-linge", level=4)
        add_info_table(doc, [
            ("Nombre", val_or_dash(sl.get("nombre"))),
            ("Capacité (kg)", val_or_dash(sl.get("capaciteKg"), " kg")),
            ("Cycles/jour", val_or_dash(data.get("cyclesSecheLinge"))),
        ])

    # Calandre
    calandre = data.get("calandre", {})
    if calandre and (calandre.get("presente") or calandre.get("present")):
        add_heading_teal(doc, "Calandre / Repasseuse", level=4)
        add_info_table(doc, [
            ("Puissance (kW)", val_or_dash(calandre.get("puissance"), " kW")),
        ])

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "buanderie", num, photos_dir)
    build_brasseur_section(doc, data)

    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "buanderie", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- CUISINE ---

CUISINE_EQUIP_DEFS = [
    {"id": "pianoGaz", "label": "Piano gaz"},
    {"id": "pianoElec", "label": "Piano électrique"},
    {"id": "pianoInduction", "label": "Piano induction"},
    {"id": "plaqueVitro", "label": "Plaque vitrocéramique"},
    {"id": "wok", "label": "Wok"},
    {"id": "grillPlancha", "label": "Grill / Plancha"},
    {"id": "sauteuseBasc", "label": "Sauteuse basculante"},
    {"id": "marmite", "label": "Marmite"},
    {"id": "fourMixte", "label": "Four mixte"},
    {"id": "fourConvection", "label": "Four à convection"},
    {"id": "fourPizza", "label": "Four à pizza"},
    {"id": "microOndes", "label": "Micro-ondes"},
    {"id": "friteuse", "label": "Friteuse"},
    {"id": "laveVaisselle", "label": "Lave-vaisselle"},
    {"id": "cfPositive", "label": "Chambre froide positive"},
    {"id": "cfNegative", "label": "Chambre froide négative"},
    {"id": "armoireRefrig", "label": "Armoire réfrigérée"},
]

def build_cuisine(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Cuisine {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Couverts/jour", val_or_dash(data.get("couvertsJour"))),
    ])

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "cuisine", num, photos_dir)
    build_brasseur_section(doc, data)

    # Équipements
    equips = data.get("equipements", {})
    if equips:
        add_heading_teal(doc, "Équipements", level=4)
        build_equipements_checkbox(doc, equips, CUISINE_EQUIP_DEFS)

    # Ventilation / Hottes
    add_info_table(doc, [
        ("Nombre de hottes", val_or_dash(data.get("nbHottes"))),
        ("Type de hottes", data.get("typeHottes", "—")),
    ])

    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "cuisine", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- BUREAUX ---

BUREAUX_EQUIP_DEFS = [
    {"id": "photocopieur", "label": "Photocopieur / multifonction"},
    {"id": "distributeur", "label": "Distributeur de boissons"},
    {"id": "fontaineEau", "label": "Fontaine à eau"},
    {"id": "machineCafe", "label": "Machine à café"},
    {"id": "microOndes", "label": "Micro-ondes"},
    {"id": "refrigerateur", "label": "Réfrigérateur"},
]

def build_bureaux(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Bureaux {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Aménagement", data.get("amenagement", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Nombre de postes", val_or_dash(data.get("nbPostes"))),
        ("Horaires", f"{data.get('horDebut', '—')} → {data.get('horFin', '—')}"),
    ])

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "bureaux", num, photos_dir)
    build_brasseur_section(doc, data)

    # Postes de travail
    add_heading_teal(doc, "Postes de travail", level=4)
    add_info_table(doc, [
        ("Type d'ordinateurs", data.get("typeOrdinateurs", "—")),
        ("Nombre d'écrans", val_or_dash(data.get("nbEcrans"))),
        ("Nombre d'imprimantes", val_or_dash(data.get("nbImprimantes"))),
    ])

    # Équipements partagés
    equips = data.get("equipPartages", {})
    if equips:
        add_heading_teal(doc, "Équipements partagés", level=4)
        build_equipements_checkbox(doc, equips, BUREAUX_EQUIP_DEFS)

    # Salle serveur
    serveur = data.get("salleServeur", {})
    if serveur and serveur.get("presente"):
        add_heading_teal(doc, "Salle serveur", level=4)
        add_info_table(doc, [
            ("Surface (m²)", val_or_dash(serveur.get("surface"), " m²")),
            ("Climatisation dédiée", oui_non(serveur.get("climDediee"))),
            ("Puissance (kW)", val_or_dash(serveur.get("puissance"), " kW")),
        ])

    # Éclairage
    add_heading_teal(doc, "Éclairage", level=4)
    add_info_table(doc, [
        ("Détection de présence", oui_non(data.get("detectionPresence", False))),
    ])
    build_eclairage_section(doc, data, "bureaux", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- SALLE DE RÉUNION ---

REUNION_AV_DEFS = [
    {"id": "videoprojecteur", "label": "Vidéoprojecteur"},
    {"id": "ecranPlat", "label": "Écran plat / TV"},
    {"id": "visioconference", "label": "Système de visioconférence"},
    {"id": "sonorisation", "label": "Sonorisation / micro"},
    {"id": "tableauInteractif", "label": "Tableau interactif / écran tactile"},
]

def build_reunion(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Salle de réunion {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Places assises", val_or_dash(data.get("placesAssises"))),
        ("Fréquence d'utilisation", data.get("frequence", "—")),
    ])

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "reunion", num, photos_dir)
    build_brasseur_section(doc, data)

    # Audiovisuel
    av = data.get("audiovisuel", {})
    if av:
        add_heading_teal(doc, "Audiovisuel", level=4)
        build_equipements_checkbox(doc, av, REUNION_AV_DEFS)

    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "reunion", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- SALLE DE SPORT ---

SPORT_EQUIP_DEFS = [
    {"id": "tapisDesCourse", "label": "Tapis de course"},
    {"id": "veloElliptique", "label": "Vélo elliptique"},
    {"id": "veloStatique", "label": "Vélo stationnaire"},
    {"id": "rameur", "label": "Rameur"},
    {"id": "stepper", "label": "Stepper / Escalier"},
    {"id": "musculation", "label": "Machines de musculation"},
    {"id": "poidsLibres", "label": "Espace poids libres"},
    {"id": "sauna", "label": "Sauna"},
    {"id": "hammam", "label": "Hammam"},
    {"id": "jacuzzi", "label": "Jacuzzi"},
]

def build_sport(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Salle de sport {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Horaires", f"{data.get('horDebut', '—')} → {data.get('horFin', '—')}"),
    ])

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "sport", num, photos_dir)
    build_brasseur_section(doc, data)

    # Équipements
    equips = data.get("equipements", {})
    if equips:
        add_heading_teal(doc, "Équipements", level=4)
        build_equipements_checkbox(doc, equips, SPORT_EQUIP_DEFS)

    # Audiovisuel
    ecrans = data.get("ecrans", {})
    if ecrans and ecrans.get("present"):
        add_heading_teal(doc, "Audiovisuel", level=4)
        add_info_table(doc, [
            ("Écrans présents", "Oui"),
            ("Nombre", val_or_dash(ecrans.get("nombre"))),
            ("Allumés en permanence", oui_non(ecrans.get("permanents"))),
            ("Sonorisation", oui_non(data.get("sono", False))),
        ])

    # Vestiaires
    vest = data.get("vestiaires", {})
    if vest and vest.get("present"):
        add_heading_teal(doc, "Vestiaires", level=4)
        add_info_table(doc, [
            ("Nombre de douches", val_or_dash(vest.get("nbDouches"))),
        ])

    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "sport", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- SALLE DE JEUX ---

JEUX_EQUIP_DEFS = [
    {"id": "billard", "label": "Billard"},
    {"id": "babyFoot", "label": "Baby-foot"},
    {"id": "pingPong", "label": "Table de ping-pong"},
    {"id": "flipper", "label": "Flipper"},
    {"id": "arcades", "label": "Bornes d'arcade"},
    {"id": "flechettes", "label": "Fléchettes"},
    {"id": "consoleJeux", "label": "Console de jeux vidéo"},
    {"id": "simulateur", "label": "Simulateur (VR, racing…)"},
    {"id": "airHockey", "label": "Air hockey"},
    {"id": "jeuxSociete", "label": "Espace jeux de société"},
    {"id": "autre", "label": "Autre"},
]

def build_jeux(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Salle de jeux {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Horaires", f"{data.get('horDebut', '—')} → {data.get('horFin', '—')}"),
    ])

    add_heading_teal(doc, "Climatisation", level=4)
    build_climatisation_section(doc, data, "jeux", num, photos_dir)
    build_brasseur_section(doc, data)

    # Équipements
    equips = data.get("equipements", {})
    if equips:
        add_heading_teal(doc, "Équipements", level=4)
        build_equipements_checkbox(doc, equips, JEUX_EQUIP_DEFS)

    # Écrans
    ecrans = data.get("ecrans", {})
    if ecrans and ecrans.get("present"):
        add_heading_teal(doc, "Écrans", level=4)
        add_info_table(doc, [
            ("Nombre", val_or_dash(ecrans.get("nombre"))),
            ("Allumés en permanence", oui_non(ecrans.get("permanents"))),
        ])

    add_heading_teal(doc, "Éclairage", level=4)
    build_eclairage_section(doc, data, "jeux", num, photos_dir)

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# --- PARKING ---

def build_parking(doc, num, data, photos_dir):
    add_heading_teal(doc, f"Parking {num} — {data.get('nom', '')}", level=3)

    add_info_table(doc, [
        ("Nom", data.get("nom", "—")),
        ("Type", data.get("type", "—")),
        ("Surface (m²)", val_or_dash(data.get("surface"), " m²")),
        ("Nombre de places", val_or_dash(data.get("nbPlaces"))),
    ])

    # Éclairage
    add_heading_teal(doc, "Éclairage", level=4)
    add_info_table(doc, [
        ("Pilotage", data.get("pilotageEclairage", "—")),
    ])
    build_eclairage_section(doc, data, "parking", num, photos_dir)

    # Ventilation
    vent = data.get("ventilation", {})
    if vent and vent.get("presente"):
        add_heading_teal(doc, "Ventilation", level=4)
        add_info_table(doc, [
            ("Puissance (kW)", val_or_dash(vent.get("puissance"), " kW")),
            ("Mode", vent.get("mode", "—")),
        ])

    # Bornes VE
    bornes = data.get("bornesVE", {})
    if bornes and bornes.get("presentes"):
        add_heading_teal(doc, "Bornes de recharge VE", level=4)
        add_info_table(doc, [
            ("Nombre", val_or_dash(bornes.get("nombre"))),
            ("Puissance unitaire (kW)", val_or_dash(bornes.get("puissance"), " kW")),
        ])

    obs = data.get("observations", "")
    if obs:
        add_heading_teal(doc, "Observations", level=4)
        add_styled_paragraph(doc, obs)

    doc.add_page_break()


# ============================================================
#   SECTION 3 — SYNTHÈSE ÉCLAIRAGE
# ============================================================

def build_synthese_eclairage(doc, audit_data):
    """Tableau récapitulatif de toute la puissance éclairage."""
    add_heading_teal(doc, "3. SYNTHÈSE — PUISSANCE ÉCLAIRAGE INSTALLÉE", level=1)

    donnees = audit_data.get("donnees", {})
    rows = []
    grand_total = 0

    module_labels = {
        "hebergements": "Hébergements", "piscines": "Piscines",
        "restaurant": "Restaurants", "bar": "Bars", "spa": "Spas",
        "buanderie": "Buanderies", "cuisine": "Cuisines",
        "bureaux": "Bureaux", "reunion": "Salles de réunion",
        "sport": "Salles de sport", "jeux": "Salles de jeux",
        "parking": "Parkings"
    }

    for mod_id, label in module_labels.items():
        mod_data = donnees.get(mod_id, {})
        mod_total = 0
        for num, data in mod_data.items():
            for ecl in data.get("eclairages", []):
                mod_total += ecl.get("puissance", 0) * ecl.get("quantite", 0)
        if mod_total > 0:
            rows.append([label, f"{mod_total} W"])
            grand_total += mod_total

    if rows:
        rows.append(["TOTAL", f"{grand_total} W"])
        add_multi_col_table(doc, ["Zone", "Puissance installée"], rows)
    else:
        add_styled_paragraph(doc, "Aucune donnée d'éclairage disponible.")

    doc.add_page_break()


# ============================================================
#   SECTION 4 — INDEX PHOTOS
# ============================================================

def build_photo_index(doc, audit_data):
    """Liste de toutes les photos de l'audit."""
    add_heading_teal(doc, "4. INDEX DES PHOTOS", level=1)

    photo_index = audit_data.get("photosIndex", [])
    if not photo_index:
        add_styled_paragraph(doc, "Aucune photo dans cet audit.")
        return

    rows = []
    for p in photo_index:
        rows.append([
            str(p.get("numero", "")),
            p.get("fichier", "—"),
            p.get("cle", "—"),
        ])

    add_multi_col_table(doc, ["N°", "Fichier", "Clé source"], rows)


# ============================================================
#   MAIN — GÉNÉRATION DU RAPPORT
# ============================================================

def generate_report(folder_path):
    """Point d'entrée principal."""

    json_path = os.path.join(folder_path, "audit.json")
    photos_dir = os.path.join(folder_path, "photos")
    output_path = os.path.join(folder_path, "rapport_audit.docx")

    if not os.path.exists(json_path):
        print(f"❌ Fichier audit.json introuvable dans : {folder_path}")
        sys.exit(1)

    print(f"📖 Lecture de {json_path}...")
    with open(json_path, "r", encoding="utf-8") as f:
        audit_data = json.load(f)

    if not os.path.isdir(photos_dir):
        photos_dir = None
        print("📷 Pas de dossier photos/ trouvé — rapport sans photos.")
    else:
        nb = len(os.listdir(photos_dir))
        print(f"📷 {nb} photo(s) trouvée(s).")

    # Création du document
    doc = Document()

    # Police par défaut
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    print("📝 Construction du rapport...")

    # 1. Page de garde
    build_cover_page(doc, audit_data)

    # 2. Modules audités
    build_modules_overview(doc, audit_data)

    # 3. État des lieux (section principale)
    build_etat_des_lieux(doc, audit_data, photos_dir)

    # 4. Synthèse éclairage
    build_synthese_eclairage(doc, audit_data)

    # 5. Index photos
    build_photo_index(doc, audit_data)

    # Sauvegarde
    doc.save(output_path)
    print(f"\n✅ Rapport généré : {output_path}")
    print(f"   ({len(audit_data.get('modulesSelectionnes', []))} modules, "
          f"{audit_data.get('meta', {}).get('nbPhotos', 0)} photos)")

    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage : python generate_report.py <dossier_rapport>")
        print("Exemple : python generate_report.py rapports/2026-05-20_14h30_Hotel")
        sys.exit(1)

    generate_report(sys.argv[1])
